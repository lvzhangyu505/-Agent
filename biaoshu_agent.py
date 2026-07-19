#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Local-first bid writing Agent for personal use.

It extracts tender requirements, builds an outline, drafts bid content from a
local knowledge base, checks compliance, and exports editable Word documents.
"""

from __future__ import annotations

import argparse
import json
import re
import textwrap
import zipfile
from dataclasses import dataclass
from datetime import datetime
from html import unescape
from pathlib import Path
from typing import Any, Dict, List, Sequence, Tuple

ROOT = Path(__file__).resolve().parent
DIRS = {
    "tenders": ROOT / "招标文件",
    "company": ROOT / "企业资料",
    "certs": ROOT / "资质证书",
    "cases": ROOT / "企业业绩",
    "history": ROOT / "历史标书",
    "templates": ROOT / "模板",
    "outputs": ROOT / "输出标书",
}

SUPPORTED = {".docx", ".doc", ".pdf", ".txt", ".md"}
DEFAULT_TEMPLATE_BOOKMARK = "ZL_INSERT_FULL_BID_BODY"

KEYWORDS = {
    "project": ["项目名称", "采购项目", "招标项目", "工程名称", "项目编号", "采购编号", "预算金额", "最高限价"],
    "qualification": ["资格要求", "供应商资格", "投标人资格", "资质", "许可证", "业绩要求", "人员要求"],
    "rejection": ["废标", "无效投标", "否决投标", "投标无效", "不予受理", "资格审查不通过"],
    "scoring": ["评分", "评审", "分值", "技术分", "商务分", "价格分", "综合评分"],
    "business": ["合同", "付款", "服务期限", "履约", "验收", "保证金", "违约", "保密"],
    "price": ["报价", "投标报价", "最高限价", "预算", "价格", "费用"],
    "format": ["暗标", "电子投标", "签字", "盖章", "页码", "目录", "封面", "格式", "装订"],
}

SECTION_TITLES = {
    "project": "项目基本信息",
    "qualification": "资格条件",
    "rejection": "废标项与否决风险",
    "scoring": "评分办法",
    "business": "商务条款",
    "price": "报价要求",
    "format": "格式与封标要求",
}


@dataclass
class SourceDoc:
    path: Path
    text: str


@dataclass
class KnowledgeChunk:
    chunk_id: str
    path: Path
    category: str
    text: str


def ensure_dirs() -> None:
    for path in DIRS.values():
        path.mkdir(parents=True, exist_ok=True)


def read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".doc":
        return (
            f"[暂不直接解析旧版 .doc：{path.name}]\n"
            "请先用 Word/WPS 另存为 .docx，或复制正文为 .txt 后再运行。"
        )
    raise ValueError(f"不支持的文件类型：{path}")


def read_docx(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        parts: List[str] = []
        parts.extend(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return read_docx_xml(path)


def read_docx_xml(path: Path) -> str:
    parts: List[str] = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml = zf.read(name).decode("utf-8", errors="ignore")
                texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
                if texts:
                    parts.append("".join(unescape(t) for t in texts))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    texts: List[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            texts.append(f"\n[第 {i} 页]\n{text}")
    return "\n".join(texts)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_sentences(text: str) -> List[str]:
    chunks = re.split(r"(?<=[。；;：:\n])", text)
    lines: List[str] = []
    for chunk in chunks:
        clean = re.sub(r"\s+", " ", chunk).strip()
        if 8 <= len(clean) <= 500:
            lines.append(clean)
    return lines


def find_keyword_lines(text: str, words: Sequence[str], limit: int = 18) -> List[str]:
    seen = set()
    matches: List[str] = []
    for line in split_sentences(text):
        if any(word in line for word in words) and line not in seen:
            seen.add(line)
            matches.append(line)
        if len(matches) >= limit:
            break
    return matches


def extract_first(patterns: Sequence[str], text: str) -> str:
    for pattern in patterns:
        m = re.search(pattern, text, re.I)
        if m:
            value = re.sub(r"\s+", " ", m.group(1)).strip(" ：:，,。；;")
            if value:
                return value[:120]
    return ""


def extract_project_meta(path: Path, text: str) -> Dict[str, str]:
    return {
        "project_name": guess_project_name(path, text),
        "project_no": extract_first([
            r"(?:项目编号|采购编号|招标编号)[:：\s]+([A-Za-z0-9\-_/（）()【】\[\]\u4e00-\u9fa5]{3,80})",
        ], text),
        "purchaser": extract_first([
            r"(?:采购人|招标人|发包人)[:：\s]+([^\n，,。；;]{3,80})",
            r"([\u4e00-\u9fa5A-Za-z0-9（）()]{3,60}(?:医院|公司|单位|中心|学校|局))",
        ], text),
        "agency": extract_first([
            r"(?:采购代理机构|招标代理机构|代理机构)[:：\s]+([^\n，,。；;]{3,80})",
        ], text),
        "budget": extract_first([
            r"(?:预算金额|采购预算|最高限价|控制价)[:：\s]*([0-9,.]+ ?(?:万元|元|人民币)?(?:/[^\n，,。；;]{1,20})?)",
        ], text),
        "service_period": extract_first([
            r"(?:服务期限|合同履行期限|履约期限|工期)[:：\s]+([^\n。；;]{2,80})",
        ], text),
        "bid_deadline": extract_first([
            r"(?:投标截止时间|响应文件提交截止时间|递交截止时间|开标时间)[:：\s]+([0-9年月日:：\-\s]{8,40})",
        ], text),
        "bid_location": extract_first([
            r"(?:开标地点|递交地点|提交地点)[:：\s]+([^\n。；;]{3,100})",
        ], text),
    }


def classify_priority(line: str, category: str) -> str:
    if category == "rejection" or any(word in line for word in ["必须", "不得", "无效", "废标", "否决", "不予受理"]):
        return "高"
    if category in {"qualification", "scoring", "format"}:
        return "中"
    return "低"


def build_requirement_items(sections: Dict[str, List[str]]) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    advice_map = {
        "project": "在投标函、封面、项目理解章节核对并统一项目基本信息。",
        "qualification": "放入资格证明文件，并在资格响应章节逐条说明。",
        "rejection": "作为封标前高风险检查项，正文和附件中必须规避。",
        "scoring": "转化为技术方案或商务证明章节，保证评审点有明确内容。",
        "business": "在商务响应表和合同条款响应中逐条确认。",
        "price": "在报价说明和报价表中保持口径、大小写金额、税费一致。",
        "format": "导出 Word 后检查页眉页脚、目录、页码、签章、暗标要求。",
    }
    for category, lines in sections.items():
        for line in lines:
            items.append({
                "category": SECTION_TITLES.get(category, category),
                "key": category,
                "priority": classify_priority(line, category),
                "source_text": line,
                "response_strategy": advice_map.get(category, "在对应章节补充响应。"),
            })
    return items


def build_scoring_items(lines: Sequence[str]) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    for line in lines:
        score = extract_first([r"([0-9]+(?:\.[0-9]+)?\s*分)"], line)
        items.append({
            "item": line[:180],
            "score": score,
            "response_chapter": guess_response_chapter(line),
            "evidence": guess_evidence(line),
        })
    return items


def build_material_items(sections: Dict[str, List[str]], structured: Dict[str, str]) -> List[Dict[str, str]]:
    items = [
        {"name": "营业执照或主体资格证明", "category": "资格材料", "source": "资格条件", "status": "待确认", "risk": "高", "advice": "核对投标主体名称、统一社会信用代码、有效期和盖章。"},
        {"name": "法定代表人身份证明/授权委托书", "category": "签章授权", "source": "格式要求", "status": "待确认", "risk": "高", "advice": "出现代理人时必须同时检查授权书、身份证明、签字和盖章。"},
        {"name": "商务条款响应表", "category": "商务响应", "source": "商务条款", "status": "待确认", "risk": "中", "advice": "服务期、付款、验收、违约责任需与招标文件一致。"},
        {"name": "报价文件/开标一览表", "category": "报价材料", "source": "报价要求", "status": "待确认", "risk": "高", "advice": "总价、分项价、大小写金额和最高限价必须闭环一致。"},
        {"name": "技术服务方案", "category": "技术响应", "source": "评分办法", "status": "待确认", "risk": "中", "advice": "按评分项组织章节，避免只有模板化描述。"},
    ]
    if sections.get("qualification"):
        items.append({"name": "资格条件逐条证明材料", "category": "资格材料", "source": sections["qualification"][0][:80], "status": "待确认", "risk": "高", "advice": "按资格条款逐项放置证明，不得只写承诺。"})
    if sections.get("rejection"):
        items.append({"name": "废标项规避证明/响应", "category": "废标风险", "source": sections["rejection"][0][:80], "status": "待确认", "risk": "高", "advice": "废标项必须逐项复核，形成封标前检查记录。"})
    if sections.get("format"):
        items.append({"name": "格式、目录、页码、签章、暗标检查", "category": "格式材料", "source": sections["format"][0][:80], "status": "待确认", "risk": "中", "advice": "导出 Word 后复核目录、页码、页眉页脚和文件属性。"})
    if structured.get("service_period"):
        items.append({"name": "服务期限承诺", "category": "商务响应", "source": structured["service_period"], "status": "待确认", "risk": "中", "advice": "所有章节中的服务期限表述需保持一致。"})
    return items


def build_timeline_items(structured: Dict[str, str], sections: Dict[str, List[str]]) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    if structured.get("bid_deadline"):
        items.append({"node": "投标截止/开标时间", "time": structured["bid_deadline"], "owner": "投标负责人", "risk": "高", "action": "至少提前半天完成最终 PDF/Word、签章、上传或密封。"})
    if structured.get("service_period"):
        items.append({"node": "服务期限", "time": structured["service_period"], "owner": "方案负责人", "risk": "中", "action": "技术方案、商务响应和承诺函保持同一服务期。"})
    for line in sections.get("business", [])[:4]:
        if any(word in line for word in ["时间", "期限", "日", "递交", "提交", "开标", "验收"]):
            items.append({"node": "商务时间节点", "time": line[:120], "owner": "商务负责人", "risk": "中", "action": "核对是否需要在响应表或承诺函中单独响应。"})
    if not items:
        items.append({"node": "封标前复核", "time": "导出 Word 后", "owner": "投标负责人", "risk": "中", "action": "检查目录、签章、页码、报价、附件和暗标风险。"})
    return items


def guess_response_chapter(line: str) -> str:
    if any(w in line for w in ["业绩", "案例", "合同"]):
        return "企业业绩"
    if any(w in line for w in ["人员", "团队", "项目负责人"]):
        return "组织架构与人员配置"
    if any(w in line for w in ["质量", "管理", "制度"]):
        return "质量控制措施"
    if any(w in line for w in ["应急", "风险"]):
        return "风险控制与应急预案"
    if any(w in line for w in ["报价", "价格"]):
        return "报价文件"
    return "技术服务方案"


def guess_evidence(line: str) -> str:
    if any(w in line for w in ["证书", "资质", "许可证"]):
        return "资质证书扫描件或证明材料"
    if any(w in line for w in ["业绩", "合同"]):
        return "类似项目合同、验收证明或中标通知书"
    if any(w in line for w in ["人员", "负责人"]):
        return "人员简历、证书、社保证明或任命文件"
    return "正文方案、承诺函或附件证明"


def guess_project_name(path: Path, text: str) -> str:
    patterns = [
        r"项目名称[:：\s]+([^\n，,。；;]{4,80})",
        r"采购项目名称[:：\s]+([^\n，,。；;]{4,80})",
        r"招标项目名称[:：\s]+([^\n，,。；;]{4,80})",
        r"([\u4e00-\u9fa5A-Za-z0-9（）()《》_-]{6,80}(?:项目|服务|采购|工程))",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return sanitize_name(m.group(1).strip())
    return sanitize_name(path.stem)


def sanitize_name(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    name = re.sub(r"\s+", "", name)
    return name[:60] or "未命名项目"


def analyze_tender(path: Path) -> Tuple[Dict, Path]:
    text = normalize_text(read_text(path))
    if not text:
        raise RuntimeError(f"未能从文件中提取文本：{path}")
    structured = extract_project_meta(path, text)
    project_name = structured["project_name"]
    out_dir = DIRS["outputs"] / project_name
    out_dir.mkdir(parents=True, exist_ok=True)

    sections = {key: find_keyword_lines(text, words) for key, words in KEYWORDS.items()}
    risks = build_risks(sections)
    checklist = build_checklist(sections)
    requirement_items = build_requirement_items(sections)
    scoring_items = build_scoring_items(sections.get("scoring", []))
    material_items = build_material_items(sections, structured)
    timeline_items = build_timeline_items(structured, sections)

    result = {
        "meta": {
            "project_name": project_name,
            "source_file": str(path),
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "text_chars": len(text),
        },
        "structured": structured,
        "sections": sections,
        "requirements": requirement_items,
        "scoring_items": scoring_items,
        "material_items": material_items,
        "timeline_items": timeline_items,
        "risks": risks,
        "checklist": checklist,
        "raw_text_sample": text[:3000],
    }
    write_json(out_dir / "01_招标文件解读.json", result)
    (out_dir / "招标文件全文.txt").write_text(text, encoding="utf-8")
    return result, out_dir


def build_risks(sections: Dict[str, List[str]]) -> List[Dict[str, str]]:
    risks: List[Dict[str, str]] = []
    if sections.get("rejection"):
        risks.append({"level": "高", "item": "存在废标/无效投标条款", "advice": "逐条建立响应清单，正文和附件中必须显式响应。"})
    if sections.get("format"):
        risks.append({"level": "中", "item": "存在格式、签章或暗标要求", "advice": "导出前检查封面、页眉页脚、公司名称、签字盖章和目录页码。"})
    if not sections.get("scoring"):
        risks.append({"level": "中", "item": "未识别到清晰评分办法", "advice": "人工确认招标文件是否另有评分表或附件。"})
    if not sections.get("qualification"):
        risks.append({"level": "中", "item": "未识别到清晰资格条件", "advice": "人工确认资格审查章节，避免漏放资质或业绩证明。"})
    return risks


def build_checklist(sections: Dict[str, List[str]]) -> List[str]:
    base = [
        "确认投标人资格条件逐条响应",
        "确认废标项/无效投标条款逐条规避",
        "确认技术方案覆盖评分点",
        "确认商务条款、服务期限、付款、验收要求已响应",
        "确认报价文件与最高限价/预算要求一致",
        "确认签字盖章、页码、目录、封面、格式符合要求",
        "确认暗标项目不出现公司名称、人员姓名、logo、历史项目识别信息",
    ]
    if sections.get("price"):
        base.append("确认报价明细、总价、税率、大小写金额一致")
    if sections.get("format"):
        base.append("导出后人工检查 Word 页眉页脚、分节、页码和目录更新")
    return base


def load_knowledge() -> List[SourceDoc]:
    docs: List[SourceDoc] = []
    for key in ["company", "certs", "cases", "history"]:
        for path in sorted(DIRS[key].rglob("*")):
            if path.is_file() and path.suffix.lower() in SUPPORTED:
                try:
                    docs.append(SourceDoc(path=path, text=normalize_text(read_text(path))))
                except Exception as exc:
                    docs.append(SourceDoc(path=path, text=f"[读取失败：{exc}]"))
    return docs


def knowledge_index_path() -> Path:
    return DIRS["company"] / "知识库分段索引.json"


def split_knowledge_text(text: str, chunk_size: int = 700, overlap: int = 100) -> List[str]:
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", text) if part.strip()]
    chunks: List[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 1 <= chunk_size:
            current = f"{current}\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        current = (current[-overlap:] + "\n" + paragraph).strip() if current else paragraph
        while len(current) > chunk_size:
            chunks.append(current[:chunk_size])
            current = current[max(1, chunk_size - overlap):]
    if current:
        chunks.append(current)
    return chunks


def build_knowledge_index() -> Dict[str, Any]:
    categories = {"company": "企业资料", "certs": "资质证书", "cases": "企业业绩", "history": "历史标书"}
    records: List[Dict[str, Any]] = []
    failures: List[Dict[str, str]] = []
    for key, label in categories.items():
        for path in sorted(DIRS[key].rglob("*")):
            if not path.is_file() or path.suffix.lower() not in SUPPORTED:
                continue
            try:
                text = normalize_text(read_text(path))
                for index, chunk in enumerate(split_knowledge_text(text), 1):
                    records.append({
                        "chunk_id": f"{key}:{path.relative_to(DIRS[key])}:{index}",
                        "path": str(path.relative_to(ROOT)),
                        "category": label,
                        "text": chunk,
                        "keywords": important_words(chunk)[:16],
                    })
            except Exception as exc:
                failures.append({"path": str(path.relative_to(ROOT)), "error": str(exc)})
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "document_count": len({item["path"] for item in records}),
        "chunk_count": len(records),
        "chunks": records,
        "failures": failures,
    }
    write_json(knowledge_index_path(), payload)
    return payload


def load_knowledge_index(rebuild: bool = False) -> Dict[str, Any]:
    path = knowledge_index_path()
    if rebuild or not path.exists():
        return build_knowledge_index()
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return build_knowledge_index()


def search_knowledge_chunks(query_words: Sequence[str], limit: int = 6) -> List[Dict[str, Any]]:
    index = load_knowledge_index()
    words = [word.strip() for word in query_words if len(word.strip()) >= 2]
    scored: List[Tuple[int, Dict[str, Any]]] = []
    for chunk in index.get("chunks", []):
        text = chunk.get("text", "")
        score = sum((4 if word in chunk.get("keywords", []) else 1) * text.count(word) for word in words)
        if score > 0:
            scored.append((score, chunk))
    scored.sort(key=lambda item: (-item[0], item[1].get("path", ""), item[1].get("chunk_id", "")))
    return [{**chunk, "score": score} for score, chunk in scored[:limit]]


def search_knowledge(query_words: Sequence[str], docs: Sequence[SourceDoc], limit: int = 6) -> List[Tuple[Path, str]]:
    scored: List[Tuple[int, Path, str]] = []
    words = [w for w in query_words if len(w) >= 2]
    for doc in docs:
        text = doc.text
        score = sum(text.count(w) for w in words)
        if score <= 0:
            continue
        excerpt = best_excerpt(text, words)
        scored.append((score, doc.path, excerpt))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [(path, excerpt) for _, path, excerpt in scored[:limit]]


def best_excerpt(text: str, words: Sequence[str], width: int = 260) -> str:
    idxs = [text.find(w) for w in words if text.find(w) >= 0]
    if not idxs:
        return text[:width]
    pos = min(idxs)
    start = max(0, pos - width // 3)
    return re.sub(r"\s+", " ", text[start : start + width]).strip()


def make_outline(analysis: Dict, out_dir: Path) -> str:
    project = analysis["meta"]["project_name"]
    sections = analysis["sections"]
    lines = [
        f"# {project} 投标文件目录",
        "",
        "## 一、投标函及资格响应",
        "1. 投标函",
        "2. 法定代表人身份证明或授权委托书",
        "3. 投标人资格证明文件",
        "4. 企业资质、证书与业绩证明",
        "",
        "## 二、商务响应文件",
        "1. 商务条款响应表",
        "2. 服务期限、付款、验收及履约承诺",
        "3. 报价说明与报价文件配合事项",
        "",
        "## 三、技术服务方案",
        "1. 项目理解与需求分析",
        "2. 服务总体方案",
        "3. 组织架构与人员配置",
        "4. 进度计划与实施保障",
        "5. 质量控制措施",
        "6. 风险控制与应急预案",
        "7. 售后/运维/持续服务方案",
        "",
        "## 四、评分点专项响应",
        "1. 评分办法逐项响应",
        "2. 关键技术指标响应",
        "3. 优势与差异化说明",
        "",
        "## 五、封标前检查清单",
    ]
    for item in analysis.get("checklist", []):
        lines.append(f"- {item}")
    if sections.get("scoring"):
        lines.extend(["", "## 识别到的评分相关原文", *[f"- {x}" for x in sections["scoring"][:8]]])
    content = "\n".join(lines) + "\n"
    (out_dir / "02_投标目录.md").write_text(content, encoding="utf-8")
    return content


def draft_bid(analysis: Dict, out_dir: Path) -> str:
    docs = load_knowledge()
    project = analysis["meta"]["project_name"]
    sections = analysis["sections"]
    company_refs = search_knowledge(["公司", "企业", "营业执照", "统一社会信用代码"], docs)
    cert_refs = search_knowledge(["资质", "证书", "许可证"], docs)
    case_refs = search_knowledge(["业绩", "合同", "项目", "服务"], docs)
    history_refs = search_knowledge(["方案", "质量", "应急", "服务", "管理"], docs)

    lines = [
        f"# {project} 投标文件初稿",
        "",
        "> 本文件由本地标书写作 Agent 生成，需结合招标原文和企业真实资料人工复核。",
        "",
        "## 一、投标函及资格响应",
        "",
        "我方已认真阅读并充分理解本项目招标文件、补充文件及相关要求，承诺按照招标文件规定提交真实、完整、有效的投标文件，并对所提供材料的真实性、准确性和合法性负责。",
        "",
        render_refs("可引用企业资料", company_refs),
        render_requirement_response("资格条件", sections.get("qualification", [])),
        render_refs("可引用资质材料", cert_refs),
        "",
        "## 二、商务响应文件",
        "",
        render_requirement_response("商务条款", sections.get("business", [])),
        render_requirement_response("报价要求", sections.get("price", [])),
        "我方承诺严格按照招标文件要求履行合同义务，接受采购人对服务质量、履约进度、成果交付和验收流程的监督管理。",
        "",
        "## 三、技术服务方案",
        "",
        "### 3.1 项目理解与需求分析",
        "根据招标文件，本项目需要围绕采购人实际业务场景、服务目标、质量要求和履约边界进行响应。我方将以合规、稳定、可追溯为原则组织实施。",
        render_requirement_response("关键响应点", sections.get("project", []) + sections.get("scoring", [])),
        "",
        "### 3.2 服务总体方案",
        "我方拟建立项目负责人统筹、专业人员分工执行、质量复核人员过程检查的服务机制，确保各项工作按计划推进、按标准交付、按要求留痕。",
        "",
        "### 3.3 组织架构与人员配置",
        "项目组织采用项目负责人制，设置实施执行、资料管理、质量复核、沟通协调等岗位。所有人员在进场或启动前完成任务交底、纪律要求和质量标准培训。",
        "",
        "### 3.4 进度计划与实施保障",
        "项目启动后，将按照准备、实施、检查、整改、验收五个阶段推进。每一阶段均设置明确成果物和责任人，重要节点形成书面记录。",
        "",
        "### 3.5 质量控制措施",
        "建立事前策划、事中检查、事后复盘的质量控制机制。对关键成果实行自检、复核、确认三级控制，发现问题及时整改并形成闭环。",
        "",
        "### 3.6 风险控制与应急预案",
        "针对人员变动、进度延误、资料缺失、沟通偏差等风险，建立备用人员、节点预警、资料台账和问题升级机制，确保服务连续性。",
        "",
        render_refs("历史标书可复用内容", history_refs),
        render_refs("可引用企业业绩", case_refs),
        "",
        "## 四、评分点专项响应",
        "",
        render_requirement_response("评分办法", sections.get("scoring", [])),
        "针对评分办法中涉及的技术方案、项目经验、人员配置、服务保障等内容，我方将在相应章节逐项展开，确保评审点有位置、有内容、有证明材料。",
        "",
        "## 五、封标前检查清单",
    ]
    for item in analysis.get("checklist", []):
        lines.append(f"- [ ] {item}")
    lines.extend(["", "## 六、待人工确认事项"])
    for risk in analysis.get("risks", []):
        lines.append(f"- 【{risk['level']}】{risk['item']}：{risk['advice']}")
    content = "\n".join(lines) + "\n"
    (out_dir / "03_标书初稿.md").write_text(content, encoding="utf-8")
    return content


def chapter_blueprint(analysis: Dict) -> List[Dict[str, Any]]:
    sections = analysis.get("sections", {})
    structured = analysis.get("structured", {})
    project = analysis.get("meta", {}).get("project_name", structured.get("project_name", "投标文件"))
    return [
        {"id": "letter", "title": "一、投标函及资格响应", "kind": "商务标", "keywords": ["公司", "企业", "资格"], "requirements": sections.get("qualification", [])},
        {"id": "business", "title": "二、商务响应文件", "kind": "商务标", "keywords": ["合同", "付款", "服务期限"], "requirements": sections.get("business", []) + sections.get("price", [])},
        {"id": "understanding", "title": "三、项目理解与需求分析", "kind": "技术标", "keywords": [project, structured.get("purchaser", ""), "服务"], "requirements": sections.get("project", [])},
        {"id": "plan", "title": "四、服务总体方案", "kind": "技术标", "keywords": ["方案", "服务", "管理"], "requirements": sections.get("scoring", [])},
        {"id": "team", "title": "五、组织架构与人员配置", "kind": "技术标", "keywords": ["人员", "团队", "项目负责人"], "requirements": find_keyword_lines("\n".join(sections.get("qualification", []) + sections.get("scoring", [])), ["人员", "团队", "负责人"], 8)},
        {"id": "quality", "title": "六、质量控制措施", "kind": "技术标", "keywords": ["质量", "检查", "考核"], "requirements": find_keyword_lines("\n".join(sections.get("scoring", []) + sections.get("business", [])), ["质量", "考核", "检查"], 8)},
        {"id": "risk", "title": "七、风险控制与应急预案", "kind": "技术标", "keywords": ["风险", "应急", "保障"], "requirements": find_keyword_lines("\n".join(sections.get("format", []) + sections.get("business", [])), ["风险", "应急", "保障"], 8)},
        {"id": "score", "title": "八、评分点专项响应", "kind": "技术标", "keywords": ["评分", "分值", "评审"], "requirements": sections.get("scoring", [])},
        {"id": "check", "title": "九、封标前检查清单", "kind": "检查", "keywords": ["检查", "暗标", "签章"], "requirements": analysis.get("checklist", [])},
    ]


def generate_chapter_content(chapter: Dict[str, Any], analysis: Dict, docs: Sequence[SourceDoc]) -> str:
    project = analysis.get("meta", {}).get("project_name", "本项目")
    structured = analysis.get("structured", {})
    query_words = list(chapter.get("keywords", []))
    for requirement in chapter.get("requirements", [])[:8]:
        query_words.extend(important_words(requirement))
    chunk_refs = search_knowledge_chunks(query_words, limit=5)
    lines = [
        f"## {chapter['title']}",
        "",
        f"本章节围绕“{chapter['title']}”进行响应，适用于{chapter.get('kind', '投标文件')}部分。投标人将结合{project}的招标要求、项目特点和履约目标，形成可执行、可检查、可追溯的响应内容。",
        "",
    ]
    if structured:
        labels = {
            "project_no": "项目编号",
            "purchaser": "采购人",
            "budget": "预算/最高限价",
            "service_period": "服务期限",
            "bid_deadline": "投标截止时间",
        }
        meta_bits = [f"{labels[k]}：{structured[k]}" for k in labels if structured.get(k)]
        if meta_bits:
            lines.extend(["### 关键信息引用", *[f"- {item}" for item in meta_bits], ""])
    requirements = chapter.get("requirements") or []
    if requirements:
        lines.append("### 招标要求逐条响应")
        for i, item in enumerate(requirements[:10], 1):
            lines.append(f"{i}. 招标要求：{item}")
            lines.append("   响应安排：我方将在本章节及相关附件中提供明确响应，确保内容、证明材料和承诺事项相互一致。")
        lines.append("")
    lines.extend([
        "### 实施与管理措施",
        "我方将建立项目负责人牵头、专业人员执行、质量复核人员检查的工作机制。对涉及进度、质量、人员、资料、沟通和验收的事项建立台账，做到任务有分工、过程有记录、问题有整改、结果可复核。",
        "",
        "### 保障与承诺",
        "我方承诺严格按照招标文件及采购人管理要求组织实施，所有响应内容均以真实资料和可执行措施为基础。若中标，将在合同签订、进场准备、服务实施、验收交付等阶段持续接受采购人监督。",
        "",
    ])
    if chunk_refs:
        lines.append("### 知识库检索与引用依据")
        for ref in chunk_refs:
            lines.append(f"- 来源：`{ref['path']}`（{ref['category']}，分段 {ref['chunk_id']}）")
            lines.append(f"  摘要：{best_excerpt(ref['text'], query_words)}")
        lines.append("")
    return "\n".join(lines)


def generate_chapters(analysis: Dict, out_dir: Path, only_id: str | None = None) -> List[Dict[str, Any]]:
    docs = load_knowledge()
    path = out_dir / "05_章节正文.json"
    existing = []
    if path.exists():
        try:
            existing = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            existing = []
    existing_map = {item.get("id"): item for item in existing}
    chapters: List[Dict[str, Any]] = []
    for item in chapter_blueprint(analysis):
        if only_id and item["id"] != only_id:
            chapters.append(existing_map.get(item["id"], item))
            continue
        query_words = list(item.get("keywords", []))
        for requirement in item.get("requirements", [])[:8]:
            query_words.extend(important_words(requirement))
        retrieved = search_knowledge_chunks(query_words, limit=5)
        content = generate_chapter_content(item, analysis, docs)
        chapters.append({
            **item,
            "content": content,
            "knowledge_refs": [{
                "chunk_id": ref["chunk_id"],
                "path": ref["path"],
                "category": ref["category"],
                "excerpt": best_excerpt(ref["text"], query_words),
                "score": ref["score"],
            } for ref in retrieved],
            "updated_at": datetime.now().isoformat(timespec="seconds"),
        })
    write_json(path, chapters)
    return chapters


def save_chapters(out_dir: Path, chapters: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    for chapter in chapters:
        chapter["updated_at"] = datetime.now().isoformat(timespec="seconds")
    write_json(out_dir / "05_章节正文.json", chapters)
    return chapters


def chapters_to_markdown(chapters: Sequence[Dict[str, Any]]) -> str:
    parts: List[str] = []
    for chapter in chapters:
        content = chapter.get("content") or ""
        if content.strip():
            parts.append(content.strip())
    return "\n\n".join(parts).strip() + "\n"


def export_chapters_docx(out_dir: Path, template: Path | None = None) -> Path:
    chapters_path = out_dir / "05_章节正文.json"
    chapters = json.loads(chapters_path.read_text(encoding="utf-8")) if chapters_path.exists() else []
    md_path = out_dir / "06_最终章节正文.md"
    md_path.write_text(chapters_to_markdown(chapters), encoding="utf-8")
    return export_docx(md_path, out_dir / "完整标书.docx", template)


REVIEW_RULES = [
    ("文件版本与范围确认", "确认招标文件版本、投标文件版本、项目/分包、采购人、投标人、截止时间和本次审查范围，防止审错版本。"),
    ("废标与资格条件优先", "先提取否决条款、资格条件、实质性要求和带特殊标记的格式，再检查对应证明材料是否存在、有效、主体一致。"),
    ("逐页逐表非空检查", "对每张表逐字段确认名称、编号、金额、数量、日期、勾选、签字、盖章是否完整；空白表不得因模板存在被视为已响应。"),
    ("全文结构与编号检查", "核对目录、章/节编号、页码、书签、正文顺序；检查错误书签、跳号、重复章节和错分包内容。"),
    ("重复内容与模板残留检查", "搜索示例文字、括号占位符、下划线、连续省略号、其他项目/医院名称、错误表名和重复表。"),
    ("横向一致性检查", "统一核对公司名称、项目名称、采购人名称、分包名称、服务期、服务地点、人员数量、报价金额、日期和承诺内容。"),
    ("报价闭环检查", "开标一览表、分项报价表、单价/数量/总价、大小写金额、最高限价和正文承诺必须闭环一致。"),
    ("授权与签章依赖检查", "出现代理人时，必须联动检查授权书、法人及代理人身份证明、签字盖章和电子签章；每份承诺函也要逐一核对。"),
    ("输出可复核问题单", "每项至少写明页码/位置、招标要求、文件现状、风险、修改动作和复核人；区分确定问题、待确认和仅建议优化。"),
]


def upgraded_review(analysis: Dict, chapters: Sequence[Dict[str, Any]], library: Dict[str, Any] | None = None) -> Dict[str, Any]:
    library = library or {}
    text = chapters_to_markdown(chapters) if chapters else ""
    structured = analysis.get("structured", {})
    sections = analysis.get("sections", {})
    material_items = analysis.get("material_items") or build_material_items(sections, structured)
    issues: List[Dict[str, str]] = []

    def add(rule: str, level: str, status: str, requirement: str, current: str, action: str, location: str = "全文件") -> None:
        issues.append({
            "rule": rule,
            "level": level,
            "status": status,
            "location": location,
            "requirement": requirement,
            "current": current,
            "risk": level,
            "action": action,
            "reviewer": "待复核",
        })

    for key, label in [("project_name", "项目名称"), ("purchaser", "采购人"), ("bid_deadline", "截止时间")]:
        value = structured.get(key)
        if value and value not in text:
            add("文件版本与范围确认", "中", "待确认", f"标书应包含并统一{label}：{value}", "正文未完整命中该字段", "在封面、投标函、项目理解和商务响应中统一该字段。")

    for line in sections.get("rejection", [])[:8]:
        if not any(word in text for word in important_words(line)[:3]):
            add("废标与资格条件优先", "高", "确定问题", line, "未在正文中找到足够明确的响应", "增加废标项规避说明，并在附件清单中确认材料存在。")

    for line in sections.get("qualification", [])[:10]:
        if not any(word in text for word in important_words(line)[:3]):
            add("废标与资格条件优先", "高", "待确认", line, "资格响应可能不充分", "补充资格证明材料，并在资格章节逐条说明。")

    placeholders = ["____", "……", "XXXX", "示例", "待填写", "【】", "（ ）", "( )", "错误！未定义书签"]
    for token in placeholders:
        if token in text:
            add("重复内容与模板残留检查", "高" if "错误" in token else "中", "确定问题", f"不得残留占位符或错误文本：{token}", f"正文出现：{token}", "定位后替换为真实内容或删除。")

    paragraphs = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if len(line.strip()) >= 40]
    duplicates: Dict[str, int] = {}
    for paragraph in paragraphs:
        key = paragraph[:180]
        duplicates[key] = duplicates.get(key, 0) + 1
    repeated = [(paragraph, count) for paragraph, count in duplicates.items() if count >= 2]
    for paragraph, count in repeated[:6]:
        add(
            "重复内容与模板残留检查", "中", "确定问题",
            "正文不应存在无必要的成段重复内容",
            f"同一段落重复 {count} 次：{paragraph[:100]}",
            "保留最符合本章节语境的一处，删除或针对章节要求重写其余重复段。",
            "章节正文",
        )

    current_names = {structured.get("project_name", ""), structured.get("purchaser", "")}
    organization_names = set(re.findall(r"[\u4e00-\u9fa5]{3,30}(?:医院|公司|中心|学校|大学|研究院|管理局)", text))
    false_positive_prefixes = ("投标人", "确认", "不得", "出现", "响应", "本项目", "招标")
    unrelated = sorted(
        name for name in organization_names
        if name
        and not any(name in current_name or current_name in name for current_name in current_names if current_name)
        and not name.startswith(false_positive_prefixes)
    )
    for name in unrelated[:6]:
        add(
            "重复内容与模板残留检查", "高", "待确认",
            "不得残留其他项目、医院或客户单位名称",
            f"正文检出可能的其他单位名称：{name}",
            "核对该名称是否为本项目合法引用；如为历史模板残留，替换或删除。",
            "全文检索",
        )

    if "报价" in text or sections.get("price"):
        has_limit = bool(structured.get("budget") or sections.get("price"))
        if has_limit and not any(word in text for word in ["大写", "小写", "总价", "单价", "分项报价"]):
            add("报价闭环检查", "高", "待确认", "报价总价、分项、大小写和限价需闭环一致", "正文报价闭环要素不足", "补齐开标一览表、分项报价表和金额一致性检查。")

    if any("授权" in c.get("title", "") or "授权" in c.get("content", "") for c in chapters):
        if not all(word in text for word in ["法定代表人", "授权委托", "身份证"]):
            add("授权与签章依赖检查", "高", "待确认", "授权链条需包含法人、代理人身份证明、签字盖章", "授权相关材料链条不完整", "补齐授权书、身份证明和签章检查。")

    if not re.search(r"(第\s*\d+\s*页|页码|目录)", text):
        add("全文结构与编号检查", "中", "仅建议优化", "目录、页码、书签、正文顺序应可复核", "正文未体现页码/目录更新状态", "导出 Word 后更新目录和页码并人工复核。")

    missing_materials = [item for item in material_items if item.get("risk") == "高"]
    for item in missing_materials[:8]:
        add("逐页逐表非空检查", item.get("risk", "中"), "待确认", item.get("name", "材料项"), "系统无法确认材料是否非空、有效、已盖章", item.get("advice", "人工打开附件逐项核对。"), "材料清单")

    for line in sections.get("rejection", [])[:12]:
        words = important_words(line)
        matched_materials = [item.get("name", "") for item in material_items if any(word in item.get("source", "") + item.get("name", "") for word in words)]
        add(
            "废标项专项检查",
            "高",
            "待确认" if matched_materials else "确定问题",
            line,
            f"关联材料：{'、'.join(matched_materials) if matched_materials else '未自动关联到证明材料'}",
            "逐字核对否决条款，确认对应表单非空、证件有效、主体一致且签章完整。",
            "废标项专项清单",
        )

    if not issues:
        add("输出可复核问题单", "低", "仅建议优化", "形成最终问题单并指定复核人", "未发现明显问题", "导出前仍需人工逐页检查签章、报价和附件。")

    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "rules": [{"name": name, "description": desc} for name, desc in REVIEW_RULES],
        "summary": {
            "total": len(issues),
            "high": sum(1 for item in issues if item["level"] == "高"),
            "medium": sum(1 for item in issues if item["level"] == "中"),
            "low": sum(1 for item in issues if item["level"] == "低"),
        },
        "issues": issues,
    }


def export_format_check(out_dir: Path, docx_path: Path | None = None) -> Dict[str, Any]:
    from docx import Document

    docx_path = docx_path or out_dir / "完整标书.docx"
    issues: List[Dict[str, str]] = []
    if not docx_path.exists() or docx_path.stat().st_size == 0:
        issues.append({"level": "高", "item": "Word 文件不存在或为空", "action": "重新导出 Word。"})
    else:
        doc = Document(str(docx_path))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs]
        nonempty = [text for text in paragraphs if text]
        if len(nonempty) < 10:
            issues.append({"level": "高", "item": f"正文段落过少，仅识别到 {len(nonempty)} 段", "action": "确认章节已生成并保存后重新导出。"})
        heading_count = sum(1 for paragraph in doc.paragraphs if paragraph.style and paragraph.style.name.startswith("Heading"))
        if heading_count == 0:
            issues.append({"level": "中", "item": "未识别到 Word 标题样式", "action": "设置章节标题样式后更新自动目录。"})
        if not any("目录" in text for text in nonempty):
            issues.append({"level": "中", "item": "未识别到目录文本", "action": "在 Word 中插入或更新自动目录。"})
        residue = [token for token in ["错误！未定义书签", "XXXX", "待填写", "____"] if any(token in text for text in nonempty)]
        for token in residue:
            issues.append({"level": "高", "item": f"导出文件仍含模板残留：{token}", "action": "定位并替换后重新导出。"})
    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "file": str(docx_path.relative_to(ROOT)) if docx_path.exists() else str(docx_path),
        "passed": not any(item["level"] == "高" for item in issues),
        "issues": issues,
    }
    write_json(out_dir / "08_导出格式检查.json", report)
    return report


def write_review_report(review: Dict[str, Any], out_dir: Path) -> str:
    lines = ["# 标书审查升级版问题单", "", f"生成时间：{review.get('generated_at', '')}", ""]
    summary = review.get("summary", {})
    lines.append(f"总问题：{summary.get('total', 0)}；高风险：{summary.get('high', 0)}；中风险：{summary.get('medium', 0)}；低风险：{summary.get('low', 0)}")
    lines.extend(["", "## 审查规则", ""])
    for rule in review.get("rules", []):
        lines.append(f"- {rule['name']}：{rule['description']}")
    lines.extend(["", "## 可复核问题单", ""])
    for i, item in enumerate(review.get("issues", []), 1):
        lines.extend([
            f"### {i}. 【{item['level']}】【{item['status']}】{item['rule']}",
            f"- 页码/位置：{item['location']}",
            f"- 招标要求：{item['requirement']}",
            f"- 文件现状：{item['current']}",
            f"- 风险：{item['risk']}",
            f"- 修改动作：{item['action']}",
            f"- 复核人：{item['reviewer']}",
            "",
        ])
    content = "\n".join(lines)
    (out_dir / "07_升级版审查问题单.md").write_text(content, encoding="utf-8")
    write_json(out_dir / "07_升级版审查问题单.json", review)
    return content


def write_library_markdown(data: Dict[str, Any]) -> None:
    company = data.get("company", {})
    legal = data.get("legal", {})
    lines = ["# 企业基础信息", ""]
    labels = {
        "company_name": "公司名称",
        "company_type": "企业类型",
        "business_period": "营业期限",
        "credit_code": "统一社会信用代码",
        "established_at": "成立时间",
        "insurance_license": "保险许可证号",
        "insurance_scope": "保险许可证业务范围",
        "regulatory_rating": "监管评级",
        "solvency_rate": "偿付能力充足率",
        "branch_count": "分支机构数量",
        "premium_income": "保费收入",
        "claim_close_rate": "理赔结案率",
    }
    for key, label in labels.items():
        value = company.get(key)
        if value:
            lines.append(f"- {label}：{value}")
    if legal:
        lines.extend(["", "## 法人信息"])
        for key, value in legal.items():
            if value:
                lines.append(f"- {key}：{value}")
    target = DIRS["company"] / "企业基础信息.md"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_refs(title: str, refs: Sequence[Tuple[Path, str]]) -> str:
    if not refs:
        return f"### {title}\n\n- 暂未在本地资料库中检索到可直接引用的材料，请补充资料后重新生成。\n"
    lines = [f"### {title}", ""]
    for path, excerpt in refs:
        lines.append(f"- 来源：`{path.relative_to(ROOT)}`")
        lines.append(f"  摘要：{excerpt}")
    return "\n".join(lines) + "\n"


def render_requirement_response(title: str, items: Sequence[str]) -> str:
    lines = [f"### {title}响应", ""]
    if not items:
        lines.append("- 招标文件中未自动识别到明确条款，请人工核对原文并补充响应。")
        return "\n".join(lines) + "\n"
    for i, item in enumerate(items[:10], start=1):
        lines.append(f"{i}. 招标要求：{item}")
        lines.append("   响应说明：我方已理解该要求，并将在投标文件对应章节提供响应说明及证明材料。")
    return "\n".join(lines) + "\n"


def compliance_check(analysis: Dict, draft_path: Path, out_dir: Path) -> str:
    draft = normalize_text(draft_path.read_text(encoding="utf-8", errors="ignore"))
    issues: List[Dict[str, str]] = []

    for key, title in SECTION_TITLES.items():
        items = analysis["sections"].get(key, [])
        for item in items[:12]:
            words = important_words(item)
            if words and not any(w in draft for w in words[:4]):
                issues.append({
                    "level": "中",
                    "area": title,
                    "problem": f"初稿可能未覆盖要求：{item[:120]}",
                    "advice": "在对应章节增加逐条响应，并补充证明材料或承诺表述。",
                })

    required_headings = ["资格", "商务", "技术", "评分", "检查"]
    for heading in required_headings:
        if heading not in draft:
            issues.append({"level": "高", "area": "目录完整性", "problem": f"缺少“{heading}”相关章节", "advice": "补齐章节后再导出。"})

    if requires_anonymous_bid(analysis["sections"].get("format", [])):
        for sensitive in ["公司名称", "我公司", "logo", "Logo", "联系人"]:
            if sensitive in draft:
                issues.append({"level": "高", "area": "暗标风险", "problem": f"初稿出现可能违反暗标的内容：{sensitive}", "advice": "暗标文件中删除或替换可识别投标人的信息。"})

    if not issues:
        issues.append({"level": "低", "area": "总体", "problem": "未发现明显缺项", "advice": "仍需人工对照招标文件、签章和最终 Word 排版。"})

    lines = [
        f"# {analysis['meta']['project_name']} 合规检查报告",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 风险问题清单",
        "",
    ]
    for i, issue in enumerate(issues, start=1):
        lines.extend([
            f"### {i}. 【{issue['level']}】{issue['area']}",
            f"- 问题：{issue['problem']}",
            f"- 建议：{issue['advice']}",
            "",
        ])
    lines.extend(["## 封标前人工复核", ""])
    for item in analysis.get("checklist", []):
        lines.append(f"- [ ] {item}")
    content = "\n".join(lines) + "\n"
    (out_dir / "04_合规检查报告.md").write_text(content, encoding="utf-8")
    return content


def important_words(text: str) -> List[str]:
    words = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,}", text)
    stop = {"项目", "招标", "投标", "文件", "要求", "采购", "应当", "必须", "不得", "进行", "提供"}
    return [w for w in words if w not in stop and len(w) >= 2][:8]


def requires_anonymous_bid(format_lines: Sequence[str]) -> bool:
    text = "\n".join(format_lines)
    if "暗标" not in text:
        return False
    negative_patterns = ["不采用暗标", "非暗标", "无需暗标", "不要求暗标", "不适用暗标", "不实行暗标"]
    if any(pattern in text for pattern in negative_patterns):
        return False
    positive_patterns = ["采用暗标", "暗标评审", "暗标编制", "按暗标", "技术标暗标", "不得出现投标人名称", "不得出现单位名称"]
    return any(pattern in text for pattern in positive_patterns)


def export_docx(md_path: Path, output: Path | None = None, template: Path | None = None) -> Path:
    from docx import Document
    from docx.shared import Pt

    output = output or md_path.with_name("完整标书.docx")
    if template and template.exists():
        doc = Document(str(template))
        doc.add_page_break()
    else:
        doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    for line in md_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        raw = line.rstrip()
        if not raw:
            doc.add_paragraph("")
            continue
        if raw.startswith("# "):
            doc.add_heading(raw[2:].strip(), level=1)
        elif raw.startswith("## "):
            doc.add_heading(raw[3:].strip(), level=2)
        elif raw.startswith("### "):
            doc.add_heading(raw[4:].strip(), level=3)
        elif raw.startswith("- [ ] "):
            doc.add_paragraph("□ " + raw[6:].strip())
        elif raw.startswith("- "):
            doc.add_paragraph(raw[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", raw):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", raw), style="List Number")
        elif raw.startswith("> "):
            doc.add_paragraph(raw[2:].strip())
        else:
            doc.add_paragraph(raw)
    doc.save(str(output))
    return output


def write_json(path: Path, data: Any) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_analysis(path: Path) -> Tuple[Dict, Path]:
    data = json.loads(path.read_text(encoding="utf-8"))
    return data, path.parent


def init_project() -> None:
    ensure_dirs()
    sample = DIRS["company"] / "企业基础信息.md"
    if not sample.exists():
        sample.write_text(
            textwrap.dedent(
                """
                # 企业基础信息

                公司名称：
                统一社会信用代码：
                企业类型：
                成立时间：
                注册地址：
                法定代表人：
                经营范围：

                ## 投标常用表述

                我方具备履行本项目所需的组织能力、人员能力、管理能力和服务保障能力。
                """
            ).strip()
            + "\n",
            encoding="utf-8",
        )
    config = ROOT / "agent_config.json"
    if not config.exists():
        write_json(config, {
            "template_bookmark": DEFAULT_TEMPLATE_BOOKMARK,
            "technical_bid_first": True,
            "provider": "local",
            "model": "",
            "api_base": "",
            "temperature": 0.3,
            "max_tokens": 4096,
        })


def find_template() -> Path | None:
    candidates = sorted(DIRS["templates"].glob("*.docx"))
    return candidates[0] if candidates else None


def run_all(tender_path: Path) -> None:
    ensure_dirs()
    analysis, out_dir = analyze_tender(tender_path)
    make_outline(analysis, out_dir)
    draft_bid(analysis, out_dir)
    generate_chapters(analysis, out_dir)
    compliance_check(analysis, out_dir / "03_标书初稿.md", out_dir)
    export_docx(out_dir / "03_标书初稿.md", out_dir / "完整标书.docx", find_template())
    print(f"已完成：{out_dir}")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="自用版标书写作 Agent")
    sub = parser.add_subparsers(dest="cmd", required=True)

    sub.add_parser("init", help="创建本地资料目录和配置模板")

    p = sub.add_parser("analyze", help="解读招标文件")
    p.add_argument("tender")

    p = sub.add_parser("outline", help="根据解读 JSON 生成目录")
    p.add_argument("analysis_json")

    p = sub.add_parser("draft", help="根据解读 JSON 和资料库生成初稿")
    p.add_argument("analysis_json")

    p = sub.add_parser("compliance", help="检查初稿合规性")
    p.add_argument("analysis_json")
    p.add_argument("draft_md")

    p = sub.add_parser("export", help="导出 Word")
    p.add_argument("draft_md")
    p.add_argument("--template", default=None)
    p.add_argument("--output", default=None)

    p = sub.add_parser("run-all", help="一键完成解读、目录、初稿、检查、Word 导出")
    p.add_argument("tender")
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    ensure_dirs()
    if args.cmd == "init":
        init_project()
        print(f"已初始化：{ROOT}")
        return 0
    if args.cmd == "analyze":
        _, out_dir = analyze_tender(Path(args.tender))
        print(f"已生成：{out_dir / '01_招标文件解读.json'}")
        return 0
    if args.cmd == "outline":
        analysis, out_dir = load_analysis(Path(args.analysis_json))
        make_outline(analysis, out_dir)
        print(f"已生成：{out_dir / '02_投标目录.md'}")
        return 0
    if args.cmd == "draft":
        analysis, out_dir = load_analysis(Path(args.analysis_json))
        draft_bid(analysis, out_dir)
        print(f"已生成：{out_dir / '03_标书初稿.md'}")
        return 0
    if args.cmd == "compliance":
        analysis, out_dir = load_analysis(Path(args.analysis_json))
        compliance_check(analysis, Path(args.draft_md), out_dir)
        print(f"已生成：{out_dir / '04_合规检查报告.md'}")
        return 0
    if args.cmd == "export":
        template = Path(args.template) if args.template else find_template()
        output = Path(args.output) if args.output else None
        result = export_docx(Path(args.draft_md), output, template)
        print(f"已生成：{result}")
        return 0
    if args.cmd == "run-all":
        run_all(Path(args.tender))
        return 0
    raise AssertionError(args.cmd)


if __name__ == "__main__":
    raise SystemExit(main())
